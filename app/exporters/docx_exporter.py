import os, json
from docx import Document
from app.file_naming import generate_filename


def _flatten_value(value) -> str:
    if isinstance(value, (dict, list)):
        return json.dumps(value, indent=2, ensure_ascii=False)
    return str(value) if value is not None else ""


def export_to_docx(data: dict, output_dir: str) -> None:
    for page in data.get("pages", [data]):
        doc = Document()
        doc_type = page.get("document_type", "Document")
        doc.add_heading(doc_type, level=1)

        for key, value in page.items():
            if key == "document_type":
                continue
            label = key.replace("_", " ").title()

            if isinstance(value, list) and value and isinstance(value[0], dict):
                doc.add_heading(label, level=2)
                headers = list(value[0].keys())
                table = doc.add_table(rows=1, cols=len(
                    headers), style="Table Grid")
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h.replace("_", " ").title()
                for row in value:
                    cells = table.add_row().cells
                    for i, h in enumerate(headers):
                        cells[i].text = _flatten_value(row.get(h))

            elif isinstance(value, dict):
                doc.add_heading(label, level=2)
                for k, v in value.items():
                    doc.add_paragraph(
                        f"{k.replace('_', ' ').title()}: {_flatten_value(v)}")

            else:
                doc.add_paragraph(f"{label}: {_flatten_value(value)}")

        path = os.path.join(output_dir, generate_filename(page, "docx"))
        doc.save(path)
        print(f"DOCX exported: {path}")
