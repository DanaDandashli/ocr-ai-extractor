from datetime import datetime
from docx import Document
from email import policy
from email.parser import BytesParser
from openpyxl import load_workbook
from bs4 import BeautifulSoup
from PIL import Image
from dotenv import load_dotenv
import xml.etree.ElementTree as ET
import os, pytesseract, fitz

load_dotenv()
pytesseract.pytesseract.tesseract_cmd = os.getenv("TESSERACT_PATH")
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"


def from_image(file_path: str, lang: str = "eng") -> str:
    text = pytesseract.image_to_string(Image.open(file_path), lang=lang)
    return text.strip()


def from_pdf(file_path: str, lang: str = "eng") -> str:
    text = ""
    with fitz.open(file_path) as doc:
        # Step 1: Try normal PDF text extraction
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
        if text.strip():
            return text.strip()
        # Step 2: OCR fallback (scanned PDF)
        for page in doc:
            pix = page.get_pixmap(colorspace=fitz.csRGB)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text += pytesseract.image_to_string(img, lang=lang) + "\n"
    return text.strip()


def pdf_to_images(pdf_path: str, output_dir: str = "temp_pages", zoom: float = 2.0) -> list[str]:
    os.makedirs(output_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    matrix = fitz.Matrix(zoom, zoom)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    paths = []
    for i in range(len(doc)):
        pix = doc.load_page(i).get_pixmap(matrix=matrix)
        path = os.path.join(output_dir, f"page_{timestamp}_{i+1}.png")
        pix.save(path)
        paths.append(path)
    return paths


def from_docx(file_path: str) -> str:
    doc = Document(file_path)
    seen = set()
    output = []
    for section in doc.sections:
        for hdr_ftr in (section.header, section.footer):
            for para in hdr_ftr.paragraphs:
                t = para.text.strip()
                if t and t not in seen:
                    seen.add(t)
                    output.append(t)
    body = doc.element.body
    for txbx in body.findall(f".//{{{W_NS}}}txbxContent"):
        parts = [t.text or "" for t in txbx.findall(f".//{{{W_NS}}}t")]
        text = "".join(parts).strip()
        if text and text not in seen:
            seen.add(text)
            output.append(text)
    for txbx in body.findall(f".//{{{WPS_NS}}}txbx"):
        parts = [t.text or "" for t in txbx.findall(f".//{{{W_NS}}}t")]
        text = "".join(parts).strip()
        if text and text not in seen:
            seen.add(text)
            output.append(text)
    for para in doc.paragraphs:
        t = para.text.strip()
        if t and t not in seen:
            seen.add(t)
            output.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in row_data:
                    row_data.append(text)
            if row_data:
                row_str = " | ".join(row_data)
                if row_str not in seen:
                    seen.add(row_str)
                    output.append(row_str)
    result = "\n".join(output).strip()
    # Prevent token overflow on very large documents
    if len(result) > 8000:
        result = result[:8000] + "\n[... truncated]"

    return result



def from_email(file_path: str) -> str:
    with open(file_path, "rb") as f:
        msg = BytesParser(policy=policy.default).parse(f)
    parts = []
    if msg["subject"]:
        parts.append(f"Subject: {msg['subject']}")
    if msg["from"]:
        parts.append(f"From: {msg['from']}")
    if msg["to"]:
        parts.append(f"To: {msg['to']}")
    body = ""
    if msg.is_multipart():
        for part in msg.walk():
            ct = part.get_content_type()
            if ct in ("text/plain", "text/html"):
                body += part.get_content()
    else:
        body = msg.get_content()
    parts.append("\nEmail Body:\n" + body)
    return "\n".join(parts).strip()


def from_excel(file_path: str) -> str:
    wb = load_workbook(file_path, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        row_count = 0
        for row in sheet.iter_rows(values_only=True):
            # Skip completely empty rows
            row_data = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if row_data:
                parts.append(" | ".join(row_data))
                row_count += 1

            # Limit rows per sheet to avoid token overflow
            if row_count >= 100:
                parts.append("[... truncated]")
                break

        parts.append("")  # blank line between sheets
    return "\n".join(parts).strip()


def from_html(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    for tag in soup(["script", "style", "noscript", "meta", "link"]):
        tag.decompose()
    return " ".join(soup.get_text(separator=" ").split()).strip()


def from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read().strip()


def from_xml(file_path: str) -> str:
    tree = ET.parse(file_path)
    root = tree.getroot()
    parts = []
    def recurse(el):
        if el.text and el.text.strip():
            parts.append(f"{el.tag}: {el.text.strip()}")
        for child in el:
            recurse(child)
        if el.tail and el.tail.strip():
            parts.append(el.tail.strip())
    recurse(root)
    return "\n".join(parts).strip()
