from docx import Document


# Namespaces
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"

NSMAP = {
    "w":   W_NS,
    "wps": WPS_NS,
}

"""
    This function extract text from .docx file (Word documents).
"""
def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    seen = set()
    output = []

    # Headers & Footers
    for section in doc.sections:
        for hdr_ftr in (section.header, section.footer):
            for para in hdr_ftr.paragraphs:
                t = para.text.strip()
                if t and t not in seen:
                    seen.add(t)
                    output.append(t)

    # Text boxes
    body = doc.element.body
    for txbx in body.findall(f".//{{{W_NS}}}txbxContent"):
        parts = []
        for t in txbx.findall(f".//{{{W_NS}}}t"):
            if t.text:
                parts.append(t.text)
        text = "".join(parts).strip()
        if text and text not in seen:
            seen.add(text)
            output.append(text)

    for txbx in body.findall(f".//{{{WPS_NS}}}txbx"):
        parts = []
        for t in txbx.findall(f".//{{{W_NS}}}t"):
            if t.text:
                parts.append(t.text)
        text = "".join(parts).strip()
        if text and text not in seen:
            seen.add(text)
            output.append(text)

    # paragraphs
    for para in doc.paragraphs:
        t = para.text.strip()
        if t and t not in seen:
            seen.add(t)
            output.append(t)

    # tables
    for table in doc.tables:
        for row in table.rows:
            row_data = []

            for cell in row.cells:
                text = cell.text.strip()

                # avoid duplicates inside same row
                if text and text not in row_data:
                    row_data.append(text)

            if row_data:
                row_str = " | ".join(row_data)

                # avoid duplicate rows
                if row_str not in seen:
                    seen.add(row_str)
                    output.append(row_str)

    return "\n".join(output).strip()
